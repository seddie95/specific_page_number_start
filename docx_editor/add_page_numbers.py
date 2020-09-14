from docx import Document
from docx.oxml import OxmlElement, ns
import os


def create_attribute(element, name, value):
    """ Create attribute for a given element """
    element.set(ns.qn(name), value)


def set_page_number_type(fmt='', start_num='1'):
    """ specify the starting page number and style"""
    # Add numbering to section
    num_type = OxmlElement('w:pgNumType')
    create_attribute(num_type, 'w:start', start_num)

    # Set the number format
    if fmt != '':
        create_attribute(num_type, 'w:fmt', fmt)

    return num_type


def add_section(parag, fmt=''):
    """Create new section tag element"""
    pPr = OxmlElement('w:pPr')
    section = OxmlElement('w:sectPr')

    # set the page number type
    nt = set_page_number_type(fmt)

    # Append the section to the paragraph
    section.append(nt)
    pPr.append(section)
    parag._p.append(pPr)


def add_page_number(parag, position=''):
    """ Add page numbers to the footer of the document"""
    run = parag.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = OxmlElement('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    # Set the position for the numbers
    if position != '':
        jc = OxmlElement('w:jc')
        create_attribute(jc, 'w:val', position)
        p = parag._p
        pPr = p.get_or_add_pPr()
        pPr.append(jc)

    # Append the new create elements to the r tag
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


# Open existing and find last paragraph before section break
doc = Document('../word_documents/steve.docx')

# Set the page number type so that it starts from one
sect = doc.sections[0]._sectPr
pgNumType = set_page_number_type()
sect.append(pgNumType)

# add numbers starting at i
last = doc.paragraphs[11]
new_paragraph = last.insert_paragraph_before()
add_section(new_paragraph, 'lowerRoman')

# Add the page numbers
add_page_number(doc.sections[0].footer.paragraphs[0],'center')

# Save a copy of the file with the page numbers
file_name = '../word_documents/steve_new.docx'
doc.save(file_name)
os.system(f'start {file_name}')
